import os
import docx
import openpyxl
from ..schemas import FileContent
from .txt_parser import get_basic_metadata

def parse_docx(file_path: str) -> FileContent:
    """
    Extracts text and metadata from a Word (.docx) file.
    """
    doc = docx.Document(file_path)
    
    # 1. Extract Text (Paragraphs + Tables)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Also grab text from tables if needed
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
                
    content = "\n".join(full_text)

    # 2. Extract Metadata (Core Properties)
    # We start with basic OS metadata...
    meta = get_basic_metadata(file_path)
    
    # ...and update it with internal Word metadata if available
    core_props = doc.core_properties
    if core_props.author:
        meta["created_author"] = core_props.author
    if core_props.last_modified_by:
        meta["modified_author"] = core_props.last_modified_by
    if core_props.created:
        meta["created_date_internal"] = core_props.created.isoformat()

    return FileContent(file_type="docx", content=content, metadata=meta)

def parse_xlsx(file_path: str) -> FileContent:
    """
    Extracts text from all sheets in an Excel (.xlsx) file.
    """
    # data_only=True gets the values, not the formulas (e.g., 50 instead of =SUM(A1:A2))
    workbook = openpyxl.load_workbook(file_path, data_only=True)
    
    full_text = []
    
    # Loop through every sheet
    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]
        full_text.append(f"--- Sheet: {sheet_name} ---")
        
        for row in sheet.iter_rows(values_only=True):
            # Filter out None values and join row data
            row_text = [str(cell) for cell in row if cell is not None]
            if row_text:
                full_text.append(" | ".join(row_text))

    content = "\n".join(full_text)

    # Metadata
    meta = get_basic_metadata(file_path)
    # Excel specifically
    if workbook.properties.creator:
        meta["created_author"] = workbook.properties.creator
    if workbook.properties.last_modified_by:
        meta["modified_author"] = workbook.properties.last_modified_by

    return FileContent(file_type="xlsx", content=content, metadata=meta)
